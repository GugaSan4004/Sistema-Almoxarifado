import os
import random
import string
import subprocess

from docx import Document
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH

class init:
    
    def __init__(self, path):
        self.path = path
        self.office = r"C:\LibreOfficePortablePrevious\App\libreoffice\program\soffice.exe"
    
    def convertPDF(self, docpath):
        output_dir = os.path.dirname(docpath)
        
        subprocess.run([
            self.office,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to", "pdf",
            docpath,
            "--outdir", output_dir
        ], check=True)

    def generate_return(self, data: dict, username: str):
        tmp_name = ''.join(random.choices((string.ascii_letters + string.digits), k=16)).upper()
                
        doc = Document(self.path / "static" / "files" / "template.docx")

        for paragraph in doc.paragraphs:
            if "ID: {{ IDCODE }}" in paragraph.text:
                paragraph.text = paragraph.text.replace("ID: {{ IDCODE }}", "")
                run = paragraph.add_run("ID: ")
                run.bold = True
                run2 = paragraph.add_run(tmp_name)
                run2.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if "ID: {{ IDCODE }}" in paragraph.text:
                            paragraph.text = paragraph.text.replace("ID: {{ IDCODE }}", "")
                            run = paragraph.add_run("ID: ")
                            run.bold = True
                            run2 = paragraph.add_run(tmp_name)
                            run2.bold = True

        table = doc.tables[0]

        rows = table.rows

        data_items = list(data.items())

        for index, (ar_code, values) in enumerate(data_items, start=1):

            if index >= len(rows):
                break

            motivo = values['reason'].upper()
            destinatario = values['name'].upper()

            cells = rows[index].cells
            
            cells[0].text = ar_code
            cells[1].text = destinatario[:39]
            cells[2].text = motivo
            
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if index == len(data_items) and index + 1 < len(rows):
                next_cells = rows[index + 1].cells

                next_cells[0].text = "***"
                next_cells[1].text = "***"
                next_cells[2].text = "***"
                
                for cell in next_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        section = doc.sections[0]
        footer = section.footer

        footer.paragraphs[0].clear()

        p = footer.paragraphs[0]
        run1 = p.add_run("DOCUMENTO EMITIDO EM:\n")
        run2 = p.add_run(f"{datetime.now().strftime('%d/%m/%Y')}\n\n")
        run3 = p.add_run("EMITIDO POR:\n")
        run4 = p.add_run(f"{username}")

        run1.bold = True
        run2.bold = True
        run3.bold = True
        run4.bold = True

        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
        docpath = self.path / "pictures" / "temp" / f"{tmp_name}.docx"
        
        doc.save(docpath)
        self.convertPDF(docpath)
        os.remove(docpath)        
        
        return tmp_name