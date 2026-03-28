import io
import os
import random
import qrcode
import string
import subprocess

from docx import Document
from docx.oxml.ns import qn 
from datetime import datetime
from docx.shared import Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE


class init:
    def __init__(self, path):
        self.path = path
        self.office = r"C:\LibreOfficePortablePrevious\App\libreoffice\program\soffice.exe"

    def convertPDF(self, docpath):
        output_dir = os.path.dirname(docpath)

        subprocess.run([
            self.office,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to", "pdf",
            docpath,
            "--outdir", output_dir
        ], check=True)

    def generate_return(self, data: dict, username: str):
        tmp_name = ''.join(random.choices(
            (string.ascii_letters + string.digits), k=16)).upper()

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        qr.add_data(tmp_name)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")

        qr_buf = io.BytesIO()
        img.save(qr_buf)
        qr_buf.seek(0)

        doc = Document(self.path / "static" / "other" / "return_template.docx")

        def make_image_floating(run, picture, width_inch, height_inch, x_inch, y_inch):
            """Makes an image floating 'In Front of Text'."""
            width = Inches(width_inch)
            height = Inches(height_inch)
            x = Inches(x_inch)
            y = Inches(y_inch)

            inline = picture._inline
            drawing = inline.getparent()
            anchor = parse_xml(f'''
                <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" {nsdecls('wp', 'a', 'pic', 'r')}>
                    <wp:simplePos x="0" y="0"/>
                    <wp:positionH relativeFrom="margin">
                        <wp:posOffset>{int(x)}</wp:posOffset>
                    </wp:positionH>
                    <wp:positionV relativeFrom="margin">
                        <wp:posOffset>{int(y)}</wp:posOffset>
                    </wp:positionV>
                    <wp:extent cx="{int(width)}" cy="{int(height)}"/>
                    <wp:effectExtent l="0" t="0" r="0" b="0"/>
                    <wp:wrapNone/>
                    <wp:docPr id="1" name="QR Code" descr="QR Code for {tmp_name}"/>
                    <wp:cNvGraphicFramePr>
                        <a:graphicFrameLocks xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" noChangeAspect="1"/>
                    </wp:cNvGraphicFramePr>
                </wp:anchor>
            ''')

            graphic = inline.getchildren()[-1]
            anchor.append(graphic)
            drawing.replace(inline, anchor)

        for paragraph in doc.paragraphs:
            if "{{ QRCODE }}" in paragraph.text:
                paragraph.text = paragraph.text = ""
                run = paragraph.add_run()
                picture = run.add_picture(qr_buf, width=Inches(0.8), height=Inches(0.8))
                make_image_floating(run, picture, 0.8, 0.8, -0.45, -0.6)

        short_name = " ".join(username.split()[:3])
        for txbx in doc.element.iter(qn('w:txbxContent')):
            for p in txbx.findall(qn('w:p')):
                runs = p.findall(f'.//{qn("w:t")}')
                full_text = "".join(r.text or "" for r in runs)
                if "{{DATE}}" in full_text or "{{NAME}}" in full_text:
                    full_text = full_text.replace("{{DATE}}", datetime.now().strftime("%d / %m / %Y"))
                    full_text = full_text.replace("{{NAME}}", short_name)
                    runs[0].text = full_text
                    for r in runs[1:]:
                        r.text = ""

        table = doc.tables[0]

        rows = table.rows

        data_items = list(data.items())

        for index, (ar_code, values) in enumerate(data_items, start=1):

            if index >= len(rows):
                break

            motivo = values['reason'].title()
            destinatario = values['name'].title()

            cells = rows[index].cells

            cells[0].text = ar_code
            cells[1].text = destinatario[:39]
            cells[2].text = motivo

            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if index == len(data_items) and index + 1 < len(rows):
                next_cells = rows[index + 1].cells

                next_cells[0].text = "***"
                next_cells[1].text = "***"
                next_cells[2].text = "***"

                for cell in next_cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in rows:
            row.height = Inches(0.25)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

        docpath = self.path / "pictures" / "temp" / f"{tmp_name}.docx"

        doc.save(docpath)
        self.convertPDF(docpath)
        os.remove(docpath)

        return tmp_name
